import json
import os
import re
import time
from io import BytesIO

import fitz  # PyMuPDF
import pandas as pd
import streamlit as st
from docx import Document
from openai import OpenAI


# ============================================================
# STREAMLIT CONFIG - 70% WIDTH ONLY
# ============================================================

st.set_page_config(
    page_title="AI-FCM Full Text Reasoning",
    page_icon="🧠",
    layout="centered"
)

st.markdown(
    """
    <style>
        .block-container {
            max-width: 70vw !important;
            padding-top: 2rem;
            padding-left: 2rem;
            padding-right: 2rem;
        }

        @media (max-width: 1200px) {
            .block-container {
                max-width: 92vw !important;
            }
        }

        textarea {
            font-size: 13px !important;
            line-height: 1.35 !important;
        }

        .stButton > button {
            width: 100%;
            border-radius: 10px;
            font-weight: 600;
        }

        .stDownloadButton > button {
            width: 100%;
            border-radius: 10px;
            font-weight: 600;
        }

        table {
            width: 100%;
        }

        th, td {
            font-size: 14px !important;
        }
    </style>
    """,
    unsafe_allow_html=True
)


# ============================================================
# AMAZON BEDROCK / MANTLE DEFAULTS
# ============================================================

DEFAULT_BEDROCK_REGION = "us-east-2"
DEFAULT_BEDROCK_MODEL = "openai.gpt-oss-120b"
DEFAULT_MAX_TOKENS = 768
DEFAULT_TEMP = 0.0
DEFAULT_MAX_DOC_CHARS = 100000
DEFAULT_REQUEST_TIMEOUT = 180.0


def build_bedrock_base_url(region: str) -> str:
    region = (region or DEFAULT_BEDROCK_REGION).strip()
    return f"https://bedrock-mantle.{region}.api.aws/v1"


def resolve_bedrock_api_key(manual_key: str = "") -> str:
    """
    Priority:
    1. Key entered in the Streamlit sidebar
    2. OPENAI_API_KEY environment variable
    3. Streamlit secret named OPENAI_API_KEY
    """
    manual_key = (manual_key or "").strip()
    if manual_key:
        return manual_key

    env_key = os.getenv("OPENAI_API_KEY", "").strip()
    if env_key:
        return env_key

    try:
        secret_key = str(st.secrets.get("OPENAI_API_KEY", "")).strip()
        if secret_key:
            return secret_key
    except Exception:
        pass

    return ""


def create_bedrock_client(api_key: str, base_url: str) -> OpenAI:
    if not api_key:
        raise ValueError(
            "Bedrock API key is missing. Enter it in the sidebar or set OPENAI_API_KEY."
        )

    return OpenAI(
        api_key=api_key,
        base_url=base_url,
        timeout=DEFAULT_REQUEST_TIMEOUT,
        max_retries=2,
    )


def test_bedrock_connection(client: OpenAI, model_id: str) -> str:
    """Send a tiny inference request so access and model permissions are truly tested."""
    response = client.chat.completions.create(
        model=model_id,
        messages=[
            {
                "role": "user",
                "content": "Reply with exactly: BEDROCK_OK",
            }
        ],
        max_completion_tokens=32,
        temperature=0.0,
    )

    return (response.choices[0].message.content or "").strip()


# ============================================================
# REQUIRED FIELDS - NEW FORMAT
# ============================================================

REQUIRED_FIELDS = [
    "Provider Phone",
    "Appointment Date",
    "Appointment Time",
    "Attorney Name",
    "Attorney Address",
    "Attorney Phone Number",
    "NCM",
    "Provider Address",
    "Provider Name (First Name / Last Name)",
    "Determining if Doctor or Provider Name",
    "Employer Name",
    "Employer Contact Name",
    "Employer Contact Email",
    "Employer Contact Mobile",
]


# ============================================================
# FIELD-ONLY PROMPT - AMAZON BEDROCK VERSION
# ============================================================

FIELD_ONLY_PROMPT = """
Extract required fields from this workers compensation / medical referral document.

Return only the final field list.
No explanation.
No markdown.
No JSON.
No thinking text.
Every field must appear once.
Never leave a field blank.
Use Not found if a real value is missing.

FIELDS:
Provider Phone:
Appointment Date:
Appointment Time:
Attorney Name:
Attorney Address:
Attorney Phone Number:
NCM:
Provider Address:
Provider Name (First Name / Last Name):
Determining if Doctor or Provider Name:
Employer Name:
Employer Contact Name:
Employer Contact Email:
Employer Contact Mobile:

RULES:

1. Provider
- Choose the MAIN provider from Provider Name, Provider / Facility Name, or provider/facility information.
- If provider text contains both a facility and a person, choose the person.
- Example: Doctors On Duty / Albert Einstein, NP -> Albert Einstein.
- Remove Dr., MD, DO, NP, PA, PA-C, FNP, PT, RN, and similar credentials from Provider Name.
- If no person exists, use the facility/clinic/hospital/practice/company name.
- Provider Phone and Provider Address must belong to the selected provider when possible.
- Do not use employer, claimant, attorney, or claims manager phone/address as provider data.
- Provider Phone must be a real phone number. If the value under Phone # is an email address, Provider Phone is Not found.
- Provider Phone must be a real phone number only.
- If the value beside Provider Phone, Phone #, or Phone Number is an email address, Provider Phone must be Not found.
- Do not place email addresses in Provider Phone.

2. Appointment
- Extract appointment date and appointment time only.
- Valid clues: Appt. Date, Appointment Date, Next Provider Appt., Provider Appt., Surgery date, NOV, or provider on date/time.
- Ignore Referral Time, Date Submitted, claim date, injury date, document date, and referral number time.
- Keep full time ranges.
- If appointment date exists but time is missing, Appointment Time is unknown time.
- If multiple different appointments exist, join dates with " & " and join times with " & ".
- If the same appointment appears twice in different formats, return it only once.
- If one source says Next Provider Appt. Date/Time and another source repeats the same date/time in Special Instructions, treat them as the same appointment.
- Do not add unknown time if a valid time exists for the same appointment date.
- NOV means Next Office Visit. Do not include NOV in Appointment Date.
- If appointment text says "NOV 09/18/2025 @ 2:30 pm", Appointment Date is 09/18/2025 and Appointment Time is 2:30 pm.
- If a line contains "@" followed by a time, use the time after @ as Appointment Time.
- Do not output unknown time when a valid time appears after @.

3. Attorney
- Use attorney section only.
- If attorney section is blank or only has labels, attorney fields are Not found.
- Do not use provider, employer, claimant, or office contact data as attorney data.

4. NCM
- NCM must be the actual nurse/NCM name.
- Valid clues: Please assign referral to, Please assign to, Customer already assigned, prior nurse, Nurse Case Manager, Tele NCM, Onsite NCM, Field NCM, assigned nurse.
- If the NCM appears as a nurse email and the username has separators, convert it to a proper name.
- Example: hilda.delgado@genexservices.com -> Hilda Delgado.
- Do not output the nurse email if a name can be inferred from the email username.
- Do not use employer/customer contact email as NCM.
- Do not use claimant, provider, attorney, or claims manager email as NCM.
- If the only email found is employer/customer contact email, NCM is Not found.
- Claims Case Manager Name is not NCM.
- Do not use Claims Case Manager Name as NCM.
- Do not use claim adjuster, claims specialist, case manager, or office contact as NCM unless the document clearly says they are the Nurse Case Manager.
- If the document only says "assigned nurse must contact claim case manager and/or tele NCM" but gives no actual nurse name/email, NCM is Not found.

5. Doctor or Facility
- Return only Doctor or Facility.
- Doctor = selected Provider Name is a real person/practitioner/specialist, including NP, PA, MD, DO, PT.
- Facility = selected Provider Name is only a clinic, hospital, center, company, office, group, or multi-doctor practice.
- If Provider Name is a person, return Doctor.

6. Employer
- Employer Name should be Customer Name, employer name, or company name.
- Do not use claimant, provider, attorney, claims manager, or employer contact person as Employer Name.

7. Employer Contact
- Use employer/customer contact information only.
- Employer Contact Name = employer/customer contact person.
- Employer Contact Email = employer/customer contact email.
- Employer Contact Mobile = employer/customer contact phone/mobile.
- Do not use claimant, provider, attorney, claims manager, or NCM phone/email.
- If no employer contact phone exists, Employer Contact Mobile is Not found.
- Employer Contact Name must not contain @ and must not be a raw username.

Employer contact name priority:
1. If a contact name is written before or near the employer contact email, use the written name.
   Example: Employer name and contact: GT Lomas-georgelomas@ups.com
   Employer Contact Name: GT Lomas
   Employer Contact Email: georgelomas@ups.com
2. If no written contact name exists, infer the name from the email username unless the username is generic.
3. Do not infer from email if a written contact name already exists.

Email-to-name fallback:
- Use only when Employer Contact Name is missing.
- Take only the username before @.
- Ignore the domain.
- Do not return the raw username.
- Do not return Not found if the username is clearly person-like.

Generic usernames are:
hr, info, payroll, benefits, claims, contact, admin, support, noreply, no-reply, customerservice, service, billing, accounting, office, mail, inbox, helpdesk, recruiting.

Separated usernames:
- robert.tester -> Robert Tester
- robert_tester -> Robert Tester
- robert-tester -> Robert Tester
- maria.santos -> Maria Santos

Combined person-like usernames:
- roberttester -> Robert Tester
- robertboden -> Robert Boden
- juelgumbs -> Juel Gumbs
- christalewing -> Christal Ewing
- mariasantos -> Maria Santos
- johnsmith -> John Smith
- georgelomas -> George Lomas

Important:
- If the document says GT Lomas-georgelomas@ups.com, use GT Lomas, not George Lomas.
- If the document only gives georgelomas@ups.com with no written name, use George Lomas.

Mandatory examples:
- roberttester@scsessidex.com means Employer Contact Name: Robert Tester
- robertboden@scsessidex.com means Employer Contact Name: Robert Boden
- juelgumbs@company.com means Employer Contact Name: Juel Gumbs

Final employer contact check:
- If Employer Contact Email is robertboden@scsessidex.com, Employer Contact Name must be Robert Boden.
- If Employer Contact Email is roberttester@scsessidex.com, Employer Contact Name must be Robert Tester.
- Employer Contact Name must not be Not found when Employer Contact Email has a non-generic person-like username.

FINAL CHECK:
- Output exactly 14 fields.
- If NCM equals Employer Contact Email and the email is employer/customer contact, set NCM to Not found.
- If Provider Name contains facility + person, keep only the person.
- If Provider Name is a person, Doctor/Facility must be Doctor.
- If Employer Contact Email exists, inspect the username before @.
- If the username is generic, Employer Contact Name is Not found.
- If the username is non-generic and person-like, Employer Contact Name must be a Proper Case full name.
- robertboden@scsessidex.com must produce Robert Boden.
- roberttester@scsessidex.com must produce Robert Tester.

DOCUMENT:
\"\"\"
{DOCUMENT_TEXT}
\"\"\"
"""


# ============================================================
# BEDROCK CLIENT
# ============================================================

# The Bedrock OpenAI-compatible client is created on demand from the
# sidebar/environment configuration. No local GGUF model is loaded.


# ============================================================
# TEXT EXTRACTION
# ============================================================

def extract_pages_from_pdf(file_bytes: bytes):
    pages = []

    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page_num, page in enumerate(doc, start=1):
            text = page.get_text("text").strip()
            pages.append({
                "page": page_num,
                "text": text
            })

    return pages


def extract_pages_from_docx(file_bytes: bytes):
    doc = Document(BytesIO(file_bytes))
    parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())

    for table_index, table in enumerate(doc.tables, start=1):
        parts.append(f"\n===== TABLE {table_index} =====")

        for row in table.rows:
            row_values = []

            for cell in row.cells:
                value = cell.text.strip().replace("\n", " ")
                row_values.append(value)

            parts.append(" | ".join(row_values))

    return [{
        "page": 1,
        "text": "\n".join(parts).strip()
    }]


def extract_pages_from_txt(file_bytes: bytes):
    try:
        text = file_bytes.decode("utf-8", errors="ignore").strip()
    except Exception:
        text = file_bytes.decode("latin-1", errors="ignore").strip()

    return [{
        "page": 1,
        "text": text
    }]


def extract_document_pages(uploaded_file):
    file_name = uploaded_file.name.lower()
    file_bytes = uploaded_file.read()

    if file_name.endswith(".pdf"):
        return extract_pages_from_pdf(file_bytes)

    if file_name.endswith(".docx"):
        return extract_pages_from_docx(file_bytes)

    if file_name.endswith(".txt"):
        return extract_pages_from_txt(file_bytes)

    raise ValueError("Unsupported file type. Upload PDF, DOCX, or TXT.")


def join_pages(pages):
    output = []

    for item in pages:
        output.append(f"\n===== PAGE {item['page']} =====\n{item['text']}")

    return "\n".join(output).strip()

def clean_joined_value(value: str) -> str:
    """
    Cleans values like:
    08/25/2025 &
    02:20PM &
    08/25/2025 & 08/25/2025
    02:20PM & 2:20pm
    """
    if not value or is_bad_value(value):
        return "Not found"

    text = str(value).strip()

    # Remove dangling separators.
    text = re.sub(r"\s*(?:&|,|;|\|)\s*$", "", text).strip()
    text = re.sub(r"^\s*(?:&|,|;|\|)\s*", "", text).strip()

    # Deduplicate joined values.
    parts = [p.strip() for p in re.split(r"\s*&\s*", text) if p.strip()]
    unique = []
    seen = set()

    for part in parts:
        key = re.sub(r"\s+", " ", part).strip().lower()

        # Normalize common time comparison: 02:20PM and 2:20pm
        key = re.sub(r"\b0([1-9]:)", r"\1", key)
        key = key.replace(" ", "")

        if key not in seen:
            seen.add(key)
            unique.append(part)

    if unique:
        return " & ".join(unique)

    return text if text else "Not found"

def fix_contact_name_acronyms(name: str) -> str:
    if not name or is_bad_value(name):
        return "Not found"

    parts = name.strip().split()
    fixed = []

    for i, part in enumerate(parts):
        clean = part.strip(".,:;")

        # Preserve/restore short initial-style names like GT, AJ, RJ, JT.
        if i == 0 and re.fullmatch(r"[A-Za-z]{2,3}", clean):
            fixed.append(clean.upper())
        else:
            fixed.append(part)

    return " ".join(fixed).strip()
# ============================================================
# OUTPUT CLEANUP / FINAL FORMAT FIX
# ============================================================

BAD_VALUES = {
    "",
    "not provided",
    "n/a",
    "none",
    "null",
    "-",
    "blank",
    "value",
    "name - phone - email",
    "address-line-1",
    "address-line-2",
    "address line 1",
    "address line 2",
    "city",
    "state",
    "zip",
    "zip code",
    "phone number",
    "mobile number",
    "cell phone",
    "email",
    "email address",
    "extension",
    "attorney information",
    "vendor information",
    "referral instructions",
    "special instructions",
    "additional instructions",
    "complete this section if applicable",
    "customer contact name",
    "customer contact phone number",
    "employer contact name",
    "employer contact phone number",
    "contact name",
    "contact phone number"
}


GENERIC_EMAIL_USERNAMES = {
    "hr",
    "info",
    "payroll",
    "benefits",
    "claims",
    "contact",
    "admin",
    "support",
    "noreply",
    "no-reply",
    "customerservice",
    "customer.service",
    "service",
    "helpdesk",
    "recruiting",
    "recruitment",
    "billing",
    "accounting",
    "office",
    "mail",
    "inbox",
}


def remove_thinking_blocks(text: str) -> str:
    text = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL | re.IGNORECASE)
    text = text.replace("<think>", "").replace("</think>", "")
    return text.strip()


def is_bad_value(value: str) -> bool:
    if value is None:
        return True

    cleaned = str(value).strip()
    cleaned_low = cleaned.lower().strip(":- ")

    if not cleaned:
        return True

    if cleaned_low.startswith("not found"):
        return True

    if "<" in cleaned or ">" in cleaned:
        return True

    if cleaned_low in BAD_VALUES:
        return True

    if cleaned_low.startswith("address-line"):
        return True

    if cleaned_low.startswith("phone number"):
        return True

    if cleaned_low.startswith("mobile number"):
        return True

    if cleaned_low.startswith("email"):
        return True

    if cleaned_low.startswith("extension"):
        return True

    if cleaned_low.startswith("referral instructions"):
        return True

    if cleaned_low.startswith("vendor information"):
        return True

    if cleaned_low.startswith("special instructions"):
        return True

    if cleaned_low.startswith("additional instructions"):
        return True

    return False


def clean_plain_text_output(text: str) -> str:
    text = remove_thinking_blocks(text)

    text = text.replace("```text", "")
    text = text.replace("```json", "")
    text = text.replace("```", "")
    text = text.replace("END", "")

    text = re.sub(r"^\s*Answer\s*:\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\n{3,}", "\n\n", text)

    match = re.search(r"Provider Phone\s*:", text, flags=re.IGNORECASE)

    if match:
        text = text[match.start():]

    lines = text.splitlines()
    kept = []
    started = False

    for line in lines:
        stripped = line.strip()

        if not stripped:
            continue

        if not started:
            if stripped.lower().startswith("provider phone:"):
                started = True
                kept.append(stripped)
            continue

        if re.match(r"(?i)^(okay|yes|that'?s|here|explanation|reasoning|note|notes|wait|let me|first)\b", stripped):
            break

        kept.append(stripped)

        if stripped.lower().startswith("employer contact mobile:"):
            break

    return "\n".join(kept).strip()

def clean_employer_contact_name(fields: dict) -> dict:
    name = fields.get("Employer Contact Name", "Not found").strip()

    bad_contact_names = {
        "customer contact name",
        "customer contact phone number",
        "employer contact name",
        "employer contact phone number",
        "contact name",
        "contact phone number",
        "extension",
        "phone number",
        "not found",
    }

    cleaned = name.lower().strip(":- ")

    if cleaned in bad_contact_names:
        fields["Employer Contact Name"] = "Not found"
        return fields

    # Reject copied labels like "Customer Contact Name:"
    if re.search(
        r"(?i)^(customer|employer)?\s*contact\s*(name|phone|number)\s*:?\s*$",
        name
    ):
        fields["Employer Contact Name"] = "Not found"
        return fields

    return fields
def parse_field_block(text: str) -> dict:
    fields = {field: "Not found" for field in REQUIRED_FIELDS}
    lines = [line.strip() for line in text.splitlines() if line.strip()]

    current_field = None

    for line in lines:
        matched_field = None

        for field in REQUIRED_FIELDS:
            prefix = field + ":"

            if line.lower().startswith(prefix.lower()):
                matched_field = field
                value = line[len(prefix):].strip()

                if value and not is_bad_value(value):
                    fields[field] = value
                else:
                    fields[field] = "Not found"

                current_field = field
                break

        if matched_field is None and current_field:
            is_another_field_label = any(
                line.lower().startswith((field + ":").lower())
                for field in REQUIRED_FIELDS
            )

            if not is_another_field_label and fields[current_field] == "Not found":
                if not is_bad_value(line):
                    fields[current_field] = line.strip()

    return fields


def extract_first_email(text: str) -> str:
    if not text:
        return ""

    match = re.search(r"[\w.\-+]+@[\w.\-]+\.\w+", str(text))
    return match.group(0).strip() if match else ""


def extract_first_phone(text: str) -> str:
    if not text:
        return ""

    match = re.search(
        r"(?<!\d)(?:\+?1[\s.\-]?)?(?:\(?\d{3}\)?[\s.\-]?)\d{3}[\s.\-]?\d{4}(?:\s*(?:x|ext|extension)\s*\d+)?(?!\d)",
        str(text),
        flags=re.IGNORECASE
    )

    return match.group(0).strip() if match else ""


def proper_case_name(name: str) -> str:
    if not name:
        return "Not found"

    parts = re.split(r"\s+", name.strip())
    cleaned = []

    for part in parts:
        if not part:
            continue

        subparts = part.split("-")
        subparts = ["'".join(x.capitalize() for x in sp.split("'")) for sp in subparts]
        cleaned.append("-".join(subparts))

    return " ".join(cleaned).strip() or "Not found"


def compact_letters(value: str) -> str:
    return re.sub(r"[^a-z]", "", str(value).lower())


def is_generic_email_username(username: str) -> bool:
    username = username.strip().lower()
    compact = compact_letters(username)

    generic_compacts = {compact_letters(x) for x in GENERIC_EMAIL_USERNAMES}
    return username in GENERIC_EMAIL_USERNAMES or compact in generic_compacts


def infer_name_from_separated_email(email: str) -> str:
    """
    Converts separated email usernames into names.
    Example:
    hilda.delgado@genexservices.com -> Hilda Delgado
    maria_santos@company.com -> Maria Santos
    john-smith@company.com -> John Smith

    Does not guess combined usernames like roberttester.
    """
    if not email or "@" not in email:
        return "Not found"

    username = email.split("@", 1)[0].strip()
    username = re.sub(r"\+.*$", "", username)

    if not username:
        return "Not found"

    if re.search(r"[._\-\s]", username):
        parts = [p for p in re.split(r"[._\-\s]+", username) if p]

        if len(parts) >= 2:
            return proper_case_name(" ".join(parts[:2]))

    return "Not found"


def clean_ncm_value(value: str) -> str:
    """
    NCM must be an actual nurse/NCM name.
    If a nurse email is provided with separators, convert it to a name.
    """
    if not value:
        return "Not found"

    v = str(value).strip()
    low = v.lower()

    invalid_phrases = [
        "assigned nurse must contact",
        "vendor nurse must contact",
        "claim case manager",
        "claims case manager",
        "liberty mutual claims specialist",
        "vendor ncm",
        "telephonic",
        "onsite",
        "full",
        "limited",
        "all communications",
        "discuss/confirm referral instructions",
        "vendor field nurse case management",
        "ncm-related instruction",
        "customer contact",
        "employer contact",
    ]

    if any(x in low for x in invalid_phrases) and "prior nurse" not in low:
        return "Not found"

    email_match = re.search(r"[\w.\-+]+@[\w.\-]+\.\w+", v)

    if email_match:
        email = email_match.group(0)
        inferred_name = infer_name_from_separated_email(email)

        if inferred_name != "Not found":
            return inferred_name

        return email

    patterns = [
        r"(?i)please\s+assign\s+referral\s+to\s+([A-Z][A-Za-z.'-]+(?:\s+[A-Z][A-Za-z.'-]+)+)",
        r"(?i)please\s+assign\s+to\s+([A-Z][A-Za-z.'-]+(?:\s+[A-Z][A-Za-z.'-]+)+)",
        r"(?i)customer\s+already\s+assigned\s+([A-Z][A-Za-z.'-]+(?:\s+[A-Z][A-Za-z.'-]+)+)\s+to\s+the\s+claim",
        r"(?i)nurse\s+case\s+manager\s*:\s*([A-Z][A-Za-z.'-]+(?:\s+[A-Z][A-Za-z.'-]+)+)",
        r"(?i)tele\s+ncm\s*:\s*([A-Z][A-Za-z.'-]+(?:\s+[A-Z][A-Za-z.'-]+)+)",
        r"(?i)onsite\s+ncm\s*:\s*([A-Z][A-Za-z.'-]+(?:\s+[A-Z][A-Za-z.'-]+)+)",
        r"(?i)field\s+ncm\s*:\s*([A-Z][A-Za-z.'-]+(?:\s+[A-Z][A-Za-z.'-]+)+)",
        r"(?i)assigned\s+nurse\s*:\s*([A-Z][A-Za-z.'-]+(?:\s+[A-Z][A-Za-z.'-]+)+)",
    ]

    for pattern in patterns:
        match = re.search(pattern, v)
        if match:
            name = match.group(1).strip()
            name = re.sub(r"\s*\([^)]*\)\s*", "", name).strip()
            return proper_case_name(name)

    if re.match(r"^[A-Z][a-zA-Z.'-]+(?:\s+[A-Z][a-zA-Z.'-]+)+$", v):
        return proper_case_name(v)

    return "Not found"

def clean_provider_name(value: str) -> str:
    if not value or is_bad_value(value):
        return "Not found"

    name = str(value).strip()

    name = re.sub(r"\s*\([^)]*priority[^)]*\)\s*", "", name, flags=re.IGNORECASE)
    name = re.sub(r"\s*\([^)]*if there is[^)]*\)\s*", "", name, flags=re.IGNORECASE)
    name = re.sub(r"\s*\([^)]*if none[^)]*\)\s*", "", name, flags=re.IGNORECASE)

    # Example: Doctors On Duty / Albert Einstein, NP -> Albert Einstein, NP
    if "/" in name:
        parts = [p.strip() for p in name.split("/") if p.strip()]
        person_parts = []

        for part in parts:
            if re.search(
                r"\b(?:Dr\.?|MD|M\.D\.|DO|D\.O\.|NP|PA|PA-C|FNP|PT|DPT|APRN|RN)\b",
                part,
                flags=re.IGNORECASE
            ):
                person_parts.append(part)

        if person_parts:
            name = person_parts[-1]
        elif len(parts) >= 2:
            name = parts[-1]

    # If the model included a facility then a person after "with" or "at".
    with_match = re.search(
        r"(?i)\b(?:with|at)\s+([A-Z][A-Za-z.'-]+(?:\s+[A-Z][A-Za-z.'-]+){1,3})(?:,\s*(?:MD|DO|NP|PA|PA-C|FNP|PT|DPT|APRN|RN))?",
        name
    )
    if with_match:
        name = with_match.group(1).strip()

    # Remove Dr. title.
    name = re.sub(r"(?i)\bDr\.?\s+", "", name)

    # Remove credentials.
    name = re.sub(
        r"(?i)\s*,?\s*\b(?:MD|M\.D\.|DO|D\.O\.|NP|PA|PA-C|FNP|PT|DPT|APRN|RN)\b\.?",
        "",
        name
    )

    # Remove accidental comments after the actual name.
    name = re.sub(r"(?i)\s*\bdoctor\s+name\s+thats\s+the\s+priority.*$", "", name)
    name = re.sub(r"(?i)\s*\bif\s+there\s+is\s+a\s+doctor.*$", "", name)

    name = re.sub(r"\s{2,}", " ", name).strip(" ,-/")

    return name if name else "Not found"


def normalize_doctor_or_facility(value: str, provider_name: str) -> str:
    value_low = (value or "").lower()
    name = (provider_name or "").strip()
    name_low = name.lower()

    facility_words = [
        "facility", "clinic", "hospital", "center", "centre", "medical center",
        "group", "company", "office", "associates", "services", "urgent care",
        "rehab", "therapy", "institute", "concentra", "rothman", "umc",
        "health", "medical", "imaging", "radiology", "doctors on duty",
        "orthopedic", "ortho", "practice", "cmc"
    ]

    practitioner_markers = [
        " md", " m.d", " do", " d.o", " np", " pa", " pa-c", " fnp",
        " pt", " dpt", " aprn", " rn", " specialist", "doctor", "dr."
    ]

    if "doctor" in value_low:
        return "Doctor"

    if any(marker in f" {name_low} " for marker in practitioner_markers):
        return "Doctor"
    
    if re.match(r"^[A-Z]{3,6}(?:\s+[A-Z][A-Za-z.'-]+){1,4}$", name):
        return "Facility"

    # Clean person-name pattern.
    if re.match(r"^[A-Z][a-zA-Z.'-]+(?:\s+[A-Z][a-zA-Z.'-]+)+$", name):
        return "Doctor"

    if len(name.split()) == 1 and not any(word in name_low for word in facility_words):
        return "Doctor"

    if any(word in name_low for word in facility_words):
        return "Facility"

    if "&" in name or " - " in name or "/" in name:
        return "Facility"

    if "facility" in value_low:
        return "Facility"

    return "Facility"


def extract_person_name_candidates(source_text: str) -> list[str]:
    """
    Extract likely employer/customer contact names from the source document.
    This avoids relying on a hardcoded first-name list.
    """
    if not source_text:
        return []

    candidates = []
    lines = [line.strip() for line in source_text.splitlines() if line.strip()]

    useful_line_keywords = [
        "customer contact",
        "employer contact",
        "contact name",
        "contact person",
        "employer name and contact",
        "customer name and contact",
        "customer contact name",
        "employer contact name",
    ]

    bad_words = {
        "customer", "employer", "contact", "name", "phone", "email",
        "address", "city", "state", "zip", "number", "extension",
        "claims", "claim", "case", "manager", "provider", "attorney",
        "vendor", "referral", "special", "instructions", "information",
        "facility", "medical", "date", "time"
    }

    for line in lines:
        low = line.lower()

        if not any(key in low for key in useful_line_keywords):
            continue

        search_area = line.split(":", 1)[1].strip() if ":" in line else line

        search_area = re.sub(r"[\w.\-+]+@[\w.\-]+\.\w+", " ", search_area)
        search_area = re.sub(
            r"(?<!\d)(?:\+?1[\s.\-]?)?(?:\(?\d{3}\)?[\s.\-]?)\d{3}[\s.\-]?\d{4}(?:\s*(?:x|ext|extension)\s*\d+)?(?!\d)",
            " ",
            search_area,
            flags=re.IGNORECASE
        )

        matches = re.findall(
            r"\b[A-Z][A-Za-z.'-]+(?:\s+[A-Z][A-Za-z.'-]+){1,3}\b",
            search_area
        )

        for match in matches:
            words = [w.strip(".,:;()[]{}") for w in match.split()]
            words_low = {w.lower() for w in words}

            if words_low & bad_words:
                continue

            cleaned = proper_case_name(" ".join(words))

            if cleaned != "Not found":
                candidates.append(cleaned)

    seen = set()
    unique = []

    for name in candidates:
        key = compact_letters(name)

        if key and key not in seen:
            seen.add(key)
            unique.append(name)

    return unique


def infer_person_name_from_email(email: str, source_text: str = "") -> str:
    """
    Safer email-to-name inference.

    Safe:
    robert.tester@company.com -> Robert Tester
    robert_tester@company.com -> Robert Tester
    robert-tester@company.com -> Robert Tester

    Conservative:
    roberttester@company.com -> Robert Tester only if the model extracted it,
    or if Robert Tester appears elsewhere in contact-related document text.

    The prompt handles combined usernames like roberttester.
    Python avoids using a hardcoded first-name list.
    """
    if not email or "@" not in email:
        return "Not found"

    username = email.split("@", 1)[0].strip().lower()
    username = re.sub(r"\+.*$", "", username)

    if not username or is_generic_email_username(username):
        return "Not found"

    # Safe case: username has separators.
    if re.search(r"[._\-\s]", username):
        parts = [p for p in re.split(r"[._\-\s]+", username) if p]

        if len(parts) >= 2:
            return proper_case_name(" ".join(parts))

        return "Not found"

    # Conservative case: combined username. Match only if a full name appears
    # in contact-related lines from the source document.
    username_compact = compact_letters(username)

    for candidate in extract_person_name_candidates(source_text):
        if compact_letters(candidate) == username_compact:
            return proper_case_name(candidate)

    return "Not found"

def infer_employer_contact_name_with_llm(
    client: OpenAI,
    model_id: str,
    email: str,
    source_text: str,
    max_tokens: int = 64,
) -> str:
    """
    Optional Bedrock correction helper:
    Employer Contact Email -> Employer Contact Name.

    This avoids hardcoded COMMON_FIRST_NAMES while still allowing:
    roberttester@scsessidex.com -> Robert Tester
    """

    if not email or "@" not in email:
        return "Not found"

    username = email.split("@", 1)[0].strip().lower()
    username = re.sub(r"\+.*$", "", username)

    if not username or is_generic_email_username(username):
        return "Not found"

    mini_prompt = f"""
You are extracting only the employer/customer contact person's name.

Employer Contact Email:
{email}

Rules:
- Return only the full person name.
- No explanation.
- No labels.
- Do not return the email.
- Do not return the raw username.
- If the username clearly represents a first name and last name, split it into Proper Case.
- roberttester means Robert Tester.
- robertboden means Robert Boden.
- juelgumbs means Juel Gumbs.
- If the username is generic like hr, info, claims, payroll, contact, admin, support, service, noreply, return Not found.
- If not confident, return Not found.

Relevant document text:
\"\"\"
{source_text[:12000]}
\"\"\"
"""

    try:
        response = client.chat.completions.create(
            model=model_id,
            messages=[
                {
                    "role": "system",
                    "content": "Return only a person name or Not found. No explanation.",
                },
                {
                    "role": "user",
                    "content": mini_prompt,
                },
            ],
            max_completion_tokens=max_tokens,
            temperature=0.0,
            top_p=0.4,
        )

        name = (response.choices[0].message.content or "").strip()
        name = re.sub(r"(?i)^Employer Contact Name\s*:\s*", "", name).strip()
        name = name.strip(" .,:;-")

        if not name or name.lower() == "not found":
            return "Not found"

        if "@" in name:
            return "Not found"

        if re.fullmatch(r"[a-z0-9._\-]+", name.strip()):
            return "Not found"

        if not re.match(r"^[A-Z][A-Za-z.'-]+(?:\s+[A-Z][A-Za-z.'-]+){1,3}$", name):
            return "Not found"

        return proper_case_name(name)

    except Exception:
        return "Not found"


def is_email_or_raw_username(value: str) -> bool:
    if not value:
        return True

    v = value.strip()

    if "@" in v:
        return True

    if re.fullmatch(r"[a-z0-9._\-]+", v):
        return True

    return False


def normalize_employer_contact_fields(fields: dict, source_text: str = "") -> dict:
    """
    Normalize the separate employer contact fields:
    Employer Contact Name
    Employer Contact Email
    Employer Contact Mobile
    """
    name = fields.get("Employer Contact Name", "Not found").strip()
    email = fields.get("Employer Contact Email", "Not found").strip()
    mobile = fields.get("Employer Contact Mobile", "Not found").strip()

    email_from_name = extract_first_email(name)
    email_from_mobile = extract_first_email(mobile)

    if is_bad_value(email):
        email = ""

    if is_bad_value(mobile):
        mobile = ""

    if email_from_name:
        email = email_from_name

    if email_from_mobile and not email:
        email = email_from_mobile

    phone_from_name = extract_first_phone(name)
    phone_from_email = extract_first_phone(email)

    if not mobile:
        if phone_from_name:
            mobile = phone_from_name
        elif phone_from_email:
            mobile = phone_from_email

    found_email = extract_first_email(email)
    email = found_email if found_email else "Not found"

    found_phone = extract_first_phone(mobile)
    mobile = found_phone if found_phone else "Not found"

    if (
        is_bad_value(name)
        or name.lower() == "not found"
        or is_email_or_raw_username(name)
    ):
        inferred_name = infer_person_name_from_email(email, source_text)
        name = inferred_name if inferred_name != "Not found" else "Not found"
    else:
        name = proper_case_name(name)

    if "@" in name or re.fullmatch(r"[a-z0-9._\-]+", name.strip()):
        inferred_name = infer_person_name_from_email(email, source_text)
        name = inferred_name if inferred_name != "Not found" else "Not found"
        
    name = fix_contact_name_acronyms(name)

    fields["Employer Contact Name"] = name
    fields["Employer Contact Email"] = email
    fields["Employer Contact Mobile"] = mobile

    return fields
def normalize_name_for_compare(value: str) -> str:
    value = str(value or "").strip()

    # Convert "Vega, Nelson" into "Nelson Vega" for comparison too.
    if "," in value:
        parts = [p.strip() for p in value.split(",", 1)]
        if len(parts) == 2:
            value = f"{parts[1]} {parts[0]}"

    return re.sub(r"[^a-z]", "", value.lower())


def protect_ncm_from_claims_case_manager(fields: dict, source_text: str = "") -> dict:
    ncm = fields.get("NCM", "Not found").strip()

    if not ncm or ncm.lower() == "not found":
        return fields

    ncm_key = normalize_name_for_compare(ncm)

    # Pull Claims Case Manager names from both same-line and next-line formats.
    patterns = [
        r"(?i)Claims\s+Case\s+Manager\s+Name\s*:\s*([^\n\r]+)",
        r"(?is)Claims\s+Case\s+Manager\s+Name\s*:\s*\n\s*([^\n\r]+)",
    ]

    for pattern in patterns:
        for match in re.findall(pattern, source_text):
            cm_name = str(match).strip()

            if not cm_name:
                continue

            if normalize_name_for_compare(cm_name) == ncm_key:
                fields["NCM"] = "Not found"
                return fields

    return fields

def clean_appointment_fields(fields: dict) -> dict:
    """
    Cleans appointment issues like:
    Appointment Date: 09/18/2025 & unknown time
    Appointment Time: 2:30 pm

    Also handles:
    NOV 09/18/2025 @ 2:30 pm
    """

    date_value = fields.get("Appointment Date", "Not found").strip()
    time_value = fields.get("Appointment Time", "Not found").strip()

    combined = f"{date_value} {time_value}"

    # Remove NOV because it means Next Office Visit, not part of the date.
    combined = re.sub(r"(?i)\bNOV\b", " ", combined)

    # Extract real date.
    date_matches = re.findall(r"\b\d{1,2}/\d{1,2}/\d{2,4}\b", combined)

    if date_matches:
        fields["Appointment Date"] = date_matches[0]
    else:
        fields["Appointment Date"] = clean_joined_value(date_value)

    # Clean time.
    time_value_clean = clean_joined_value(time_value)

    # If time is missing/unknown, try extracting from combined text.
    if (
        not time_value_clean
        or time_value_clean.lower() == "not found"
        or "unknown time" in time_value_clean.lower()
    ):
        found_time = extract_time_from_text(combined)
        fields["Appointment Time"] = found_time if found_time else "unknown time"
    else:
        fields["Appointment Time"] = time_value_clean

    # Remove dangling ampersand if still present.
    fields["Appointment Date"] = re.sub(r"\s*&\s*$", "", fields["Appointment Date"]).strip()
    fields["Appointment Time"] = re.sub(r"\s*&\s*$", "", fields["Appointment Time"]).strip()

    # Do not allow "unknown time" inside Appointment Date.
    fields["Appointment Date"] = re.sub(
        r"(?i)\s*&?\s*unknown\s+time\s*",
        "",
        fields["Appointment Date"]
    ).strip()

    if fields["Appointment Date"] != "Not found" and fields["Appointment Time"] == "Not found":
        fields["Appointment Time"] = "unknown time"

        fields["Appointment Date"] = add_date_fix_suggestion(
        fields.get("Appointment Date", "")
    )
    return fields

def protect_ncm_from_employer_contact(fields: dict) -> dict:
    ncm = fields.get("NCM", "Not found").strip()
    employer_email = fields.get("Employer Contact Email", "Not found").strip()

    ncm_email = extract_first_email(ncm)
    emp_email = extract_first_email(employer_email)

    if ncm_email and emp_email and ncm_email.lower() == emp_email.lower():
        fields["NCM"] = "Not found"

    return fields


def extract_time_from_text(text: str) -> str:
    if not text:
        return ""

    time_pattern = r"""
        \b
        (?:
            (?:1[0-2]|0?[1-9])
            (?::[0-5]\d)?
            \s*
            (?:a\.?m\.?|p\.?m\.?|am|pm|AM|PM)
            |
            (?:[01]?\d|2[0-3])
            :
            [0-5]\d
        )
        (?:\s*(?:PST|EST|CST|MST|PT|ET|CT|MT))?
        (?:\s*[-–]\s*
            (?:
                (?:1[0-2]|0?[1-9])
                (?::[0-5]\d)?
                \s*
                (?:a\.?m\.?|p\.?m\.?|am|pm|AM|PM)
                |
                (?:[01]?\d|2[0-3])
                :
                [0-5]\d
            )
            (?:\s*(?:PST|EST|CST|MST|PT|ET|CT|MT))?
        )?
        \b
    """

    match = re.search(time_pattern, text, flags=re.IGNORECASE | re.VERBOSE)
    return match.group(0).strip() if match else ""

def add_date_fix_suggestion(date_value: str) -> str:
    """
    Keeps the original appointment date but adds a correction hint
    when the year looks like an obvious typo.

    Example:
    08/22/2205 -> 08/22/2205 (fix: 08/22/2025)
    """

    if not date_value or date_value.lower() == "not found":
        return date_value

    text = str(date_value).strip()

    match = re.search(r"\b(\d{1,2})/(\d{1,2})/(\d{4})\b", text)

    if not match:
        return text

    month, day, year = match.groups()
    year_int = int(year)

    # Common referral year typo: 2205 instead of 2025
    if 2100 <= year_int <= 2299:
        fixed_year = "20" + year[-2:]
        fixed_date = f"{month}/{day}/{fixed_year}"

        if fixed_date not in text:
            return f"{text} (fix: {fixed_date})"

    return text

def split_appointment_date_time(fields: dict) -> dict:
    """
    Safety helper in case the LLM puts date and time together.
    """
    date_value = fields.get("Appointment Date", "Not found").strip()
    time_value = fields.get("Appointment Time", "Not found").strip()

    if is_bad_value(date_value):
        fields["Appointment Date"] = "Not found"

    if is_bad_value(time_value):
        fields["Appointment Time"] = "Not found"

    date_value = fields.get("Appointment Date", "Not found").strip()
    time_value = fields.get("Appointment Time", "Not found").strip()

    if time_value.lower() == "not found":
        found_time = extract_time_from_text(date_value)

        if found_time:
            fields["Appointment Time"] = found_time
            fields["Appointment Date"] = date_value.replace(found_time, "").strip(" -–,|")

    if fields.get("Appointment Time", "").strip().lower() == "not found":
        if fields.get("Appointment Date", "").strip().lower() != "not found":
            fields["Appointment Time"] = "unknown time"

    fields["Appointment Date"] = clean_joined_value(fields.get("Appointment Date", ""))
    fields["Appointment Time"] = clean_joined_value(fields.get("Appointment Time", ""))

    if fields["Appointment Date"] != "Not found" and fields["Appointment Time"] == "Not found":
        fields["Appointment Time"] = "unknown time"

    return fields


def format_field_block(fields: dict) -> str:
    output = []

    for field in REQUIRED_FIELDS:
        value = fields.get(field, "Not found")

        if is_bad_value(value):
            value = "Not found"

        output.append(f"{field}: {value}")

    return "\n".join(output)


def fields_to_table_df(fields: dict) -> pd.DataFrame:
    """
    Clean table format:
    Provider Phone        value
    Appointment Date      value
    """
    rows = []

    for field in REQUIRED_FIELDS:
        value = fields.get(field, "Not found")

        if is_bad_value(value):
            value = "Not found"

        rows.append({
            "Field": field,
            "Value": value
        })

    return pd.DataFrame(rows)


def table_df_to_csv_text(df: pd.DataFrame) -> str:
    return df.to_csv(index=False)

def phone_or_not_found(value: str) -> str:
    phone = extract_first_phone(value)
    return phone if phone else "Not found"


def clean_provider_phone(fields: dict) -> dict:
    fields["Provider Phone"] = phone_or_not_found(fields.get("Provider Phone", ""))
    return fields

def force_exact_field_output(text: str, source_text: str = "") -> tuple[str, dict]:
    cleaned = clean_plain_text_output(text)
    fields = parse_field_block(cleaned)

    fields["NCM"] = clean_ncm_value(fields.get("NCM", ""))

    fields["Provider Name (First Name / Last Name)"] = clean_provider_name(
        fields.get("Provider Name (First Name / Last Name)", "")
    )

    provider_name = fields.get("Provider Name (First Name / Last Name)", "")
    determine_value = fields.get("Determining if Doctor or Provider Name", "")

    fields["Determining if Doctor or Provider Name"] = normalize_doctor_or_facility(
        determine_value,
        provider_name
    )

    for field in ["Attorney Name", "Attorney Address", "Attorney Phone Number"]:
        if is_bad_value(fields.get(field, "")):
            fields[field] = "Not found"

    fields = split_appointment_date_time(fields)
    fields = clean_appointment_fields(fields)

    fields = normalize_employer_contact_fields(fields, source_text)
    fields = clean_employer_contact_name(fields)

    fields = protect_ncm_from_employer_contact(fields)
    fields = protect_ncm_from_claims_case_manager(fields, source_text)
    fields = clean_provider_phone(fields)

    final_text = format_field_block(fields)
    return final_text, fields

def too_many_missing_fields(result_text: str) -> bool:
    fields = parse_field_block(result_text)

    missing = 0

    for _, value in fields.items():
        if is_bad_value(value) or value.strip().lower() == "not found":
            missing += 1

    return missing >= 9


# ============================================================
# AMAZON BEDROCK REASONING
# ============================================================

def call_llm_once(
    client: OpenAI,
    model_id: str,
    prompt: str,
    max_tokens: int,
    temperature: float,
):
    response = client.chat.completions.create(
        model=model_id,
        messages=[
            {
                "role": "system",
                "content": (
                    "You are a strict document field extraction engine. "
                    "Return only the requested completed field list. "
                    "Do not explain. Do not provide reasoning. "
                    "Do not copy blank templates."
                ),
            },
            {
                "role": "user",
                "content": prompt,
            },
        ],
        max_completion_tokens=max_tokens,
        temperature=temperature,
        top_p=0.4,
    )

    return (response.choices[0].message.content or "").strip()


def run_reasoning(
    client: OpenAI,
    model_id: str,
    full_document_text: str,
    validation_source_text: str = "",
    max_tokens: int = 768,
    temperature: float = 0.0,
):
    """
    full_document_text:
        Extracted text sent to Amazon Bedrock.

    validation_source_text:
        Original full extracted text used by Python validators.
    """

    if not validation_source_text:
        validation_source_text = full_document_text

    prompt = FIELD_ONLY_PROMPT.replace("{DOCUMENT_TEXT}", full_document_text)

    start_time = time.time()

    output_text = call_llm_once(
        client=client,
        model_id=model_id,
        prompt=prompt,
        max_tokens=max_tokens,
        temperature=temperature,
    )

    cleaned = clean_plain_text_output(output_text)

    final_result, final_fields = force_exact_field_output(
        cleaned,
        validation_source_text,
    )

    if too_many_missing_fields(final_result):
        retry_prompt = prompt + """

IMPORTANT CORRECTION:
Your previous output was mostly blank or copied the empty field list.
Do not copy the empty field list.
Read the document again and fill the actual values.
Return only the 14 completed lines.
"""

        output_text = call_llm_once(
            client=client,
            model_id=model_id,
            prompt=retry_prompt,
            max_tokens=max_tokens,
            temperature=temperature,
        )

        cleaned = clean_plain_text_output(output_text)

        final_result, final_fields = force_exact_field_output(
            cleaned,
            validation_source_text,
        )

    elapsed = round(time.time() - start_time, 2)

    return final_result, final_fields, cleaned, elapsed


# ============================================================
# UI
# ============================================================

st.title("🧠 AI-FCM Bedrock Extraction")
st.caption(
    "Selectable PDF / DOCX / TXT → extracted text → Amazon Bedrock "
    "(OpenAI-compatible Mantle API) → validated required fields"
)

with st.sidebar:
    st.header("Amazon Bedrock Settings")

    bedrock_region = st.text_input(
        "AWS Region",
        value=os.getenv("AWS_REGION", DEFAULT_BEDROCK_REGION),
        help="Your screenshot shows US East (Ohio), which is us-east-2.",
    ).strip()

    bedrock_model = st.text_input(
        "Model ID",
        value=os.getenv("BEDROCK_MODEL_ID", DEFAULT_BEDROCK_MODEL),
        help="For the Mantle endpoint, gpt-oss-120b uses openai.gpt-oss-120b.",
    ).strip()

    manual_api_key = st.text_input(
        "Bedrock API key",
        value="",
        type="password",
        placeholder="Leave blank to use OPENAI_API_KEY",
        help=(
            "For production, prefer OPENAI_API_KEY in the environment or "
            "Streamlit secrets instead of saving the key in source code."
        ),
    )

    bedrock_api_key = resolve_bedrock_api_key(manual_api_key)
    bedrock_base_url = build_bedrock_base_url(bedrock_region)

    st.caption("Mantle endpoint")
    st.code(bedrock_base_url, language=None)

    if bedrock_api_key:
        if manual_api_key.strip():
            st.success("Bedrock API key entered for this session.")
        else:
            st.success("Bedrock API key found in environment/secrets.")
    else:
        st.warning("No Bedrock API key configured yet.")

    if st.button("Test Bedrock Connection"):
        if not bedrock_region:
            st.error("Enter an AWS region.")
        elif not bedrock_model:
            st.error("Enter a Bedrock model ID.")
        elif not bedrock_api_key:
            st.error("Enter a Bedrock API key or set OPENAI_API_KEY.")
        else:
            try:
                with st.spinner("Testing Bedrock inference access..."):
                    test_client = create_bedrock_client(
                        api_key=bedrock_api_key,
                        base_url=bedrock_base_url,
                    )
                    test_result = test_bedrock_connection(test_client, bedrock_model)

                if "BEDROCK_OK" in test_result.upper():
                    st.success("Bedrock connection and inference are working.")
                else:
                    st.success("Bedrock responded successfully.")
                    st.caption(f"Response: {test_result[:200] or '[empty content]'}")
            except Exception as e:
                st.error(f"Bedrock test failed: {e}")

    st.divider()
    st.subheader("Inference Settings")

    max_tokens = st.number_input(
        "Max output tokens",
        min_value=128,
        max_value=4096,
        value=DEFAULT_MAX_TOKENS,
        step=128,
    )

    temperature = st.slider(
        "Temperature",
        min_value=0.0,
        max_value=0.5,
        value=DEFAULT_TEMP,
        step=0.05,
    )

    max_doc_chars = st.number_input(
        "Max document characters sent to Bedrock",
        min_value=5000,
        max_value=400000,
        value=DEFAULT_MAX_DOC_CHARS,
        step=5000,
        help=(
            "gpt-oss-120b has a large context window, but keeping a cap avoids "
            "accidentally sending extremely large documents."
        ),
    )

uploaded_file = st.file_uploader(
    "Upload selectable PDF, DOCX, or TXT",
    type=["pdf", "docx", "txt"],
)

if uploaded_file:
    st.subheader("1. Extract Full Document Text")

    if st.button("Extract Full Text"):
        try:
            status = st.empty()

            status.info("Reading document...")
            pages = extract_document_pages(uploaded_file)

            status.info("Building full extracted text...")
            full_text = join_pages(pages)

            text_for_bedrock = full_text

            if len(text_for_bedrock) > int(max_doc_chars):
                text_for_bedrock = text_for_bedrock[: int(max_doc_chars)]
                status.warning(
                    f"Full text has {len(full_text)} characters. "
                    f"Sending first {int(max_doc_chars)} characters based on current setting."
                )
            else:
                status.success(
                    f"Full text extracted. Sending all {len(full_text)} characters to Bedrock."
                )

            st.session_state["pages"] = pages
            st.session_state["full_text"] = full_text
            st.session_state["text_for_bedrock"] = text_for_bedrock

        except Exception as e:
            st.error(f"Extraction failed: {e}")

if "full_text" in st.session_state:
    st.divider()

    col1, col2 = st.columns([1, 1])

    with col1:
        st.subheader("Full Extracted Text")
        st.write(f"Characters: **{len(st.session_state['full_text'])}**")

        with st.expander("View full extracted text", expanded=False):
            st.text_area(
                "Full extracted text",
                value=st.session_state["full_text"],
                height=350,
            )

        st.download_button(
            "Download full extracted text",
            data=st.session_state["full_text"],
            file_name="full_extracted_text.txt",
            mime="text/plain",
        )

    with col2:
        text_for_bedrock = st.session_state.get(
            "text_for_bedrock",
            st.session_state.get("text_for_llm", ""),
        )

        st.subheader("Text Sent to Bedrock")
        st.write(f"Characters sent: **{len(text_for_bedrock)}**")

        with st.expander("View text sent to Bedrock", expanded=True):
            st.text_area(
                "Text sent to Bedrock",
                value=text_for_bedrock,
                height=350,
            )

        st.download_button(
            "Download text sent to Bedrock",
            data=text_for_bedrock,
            file_name="text_sent_to_bedrock.txt",
            mime="text/plain",
        )

    st.divider()
    st.subheader("2. Run Amazon Bedrock Extraction")

    if st.button("Extract Required Fields"):
        if not bedrock_region:
            st.error("Please enter the AWS region.")
        elif not bedrock_model:
            st.error("Please enter the Bedrock model ID.")
        elif not bedrock_api_key:
            st.error(
                "Bedrock API key is missing. Enter it in the sidebar or set OPENAI_API_KEY."
            )
        elif not text_for_bedrock.strip():
            st.error("No document text available.")
        else:
            try:
                status = st.empty()

                status.info(f"Connecting to Amazon Bedrock: {bedrock_model}...")
                client = create_bedrock_client(
                    api_key=bedrock_api_key,
                    base_url=bedrock_base_url,
                )

                status.info("Sending extracted document text to Amazon Bedrock...")
                final_result, final_fields, raw_cleaned, elapsed = run_reasoning(
                    client=client,
                    model_id=bedrock_model,
                    full_document_text=text_for_bedrock,
                    validation_source_text=st.session_state["full_text"],
                    max_tokens=int(max_tokens),
                    temperature=float(temperature),
                )

                result_df = fields_to_table_df(final_fields)
                result_json = json.dumps(final_fields, indent=2, ensure_ascii=False)

                st.session_state["result_text"] = final_result
                st.session_state["result_fields"] = final_fields
                st.session_state["result_df"] = result_df
                st.session_state["result_json"] = result_json
                st.session_state["raw_cleaned"] = raw_cleaned
                st.session_state["elapsed"] = elapsed
                st.session_state["bedrock_model_used"] = bedrock_model

                status.success(f"Bedrock extraction complete in {elapsed} seconds.")

            except Exception as e:
                st.error(f"Amazon Bedrock request failed: {e}")

if "result_df" in st.session_state:
    st.divider()
    st.subheader("Required Fields Result")

    if "elapsed" in st.session_state:
        model_used = st.session_state.get("bedrock_model_used", bedrock_model)
        st.info(
            f"Bedrock model: {model_used} | "
            f"request/reasoning time: {st.session_state['elapsed']} seconds"
        )

    result_df = st.session_state["result_df"]
    result_text = st.session_state["result_text"]
    result_json = st.session_state.get(
        "result_json",
        json.dumps(st.session_state.get("result_fields", {}), indent=2),
    )

    st.dataframe(
        result_df,
        hide_index=True,
        use_container_width=True,
    )

    csv_text = table_df_to_csv_text(result_df)

    col_csv, col_json, col_txt = st.columns(3)

    with col_csv:
        st.download_button(
            "Download CSV",
            data=csv_text,
            file_name="ai_fcm_required_fields_table.csv",
            mime="text/csv",
        )

    with col_json:
        st.download_button(
            "Download JSON",
            data=result_json,
            file_name="ai_fcm_required_fields.json",
            mime="application/json",
        )

    with col_txt:
        st.download_button(
            "Download TXT",
            data=result_text,
            file_name="ai_fcm_required_fields.txt",
            mime="text/plain",
        )

    with st.expander("JSON output for RPA", expanded=True):
        st.code(result_json, language="json")

    with st.expander("Plain text output", expanded=False):
        st.text_area(
            "Final output",
            value=result_text,
            height=350,
        )

    with st.expander("Raw cleaned Bedrock output", expanded=False):
        st.text_area(
            "Raw cleaned output",
            value=st.session_state.get("raw_cleaned", ""),
            height=250,
        )
